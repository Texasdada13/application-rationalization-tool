#!/usr/bin/env python3
"""
Convert Markdown files to Word (.docx) format
"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_docx(md_file, docx_file):
    """Convert a markdown file to a Word document"""

    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Process line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                # End of code block
                in_code_block = False
                if code_block_lines:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    for code_line in code_block_lines:
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)

        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')

        # Handle horizontal rules
        elif line.strip() == '---' or line.strip() == '***':
            doc.add_paragraph('_' * 50)

        # Handle empty lines
        elif line.strip() == '':
            doc.add_paragraph()

        # Regular text with inline formatting
        else:
            p = doc.add_paragraph()
            # Simple inline formatting
            text = line

            # Bold and italic handling
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
            for part in parts:
                if part.startswith('***') and part.endswith('***'):
                    run = p.add_run(part[3:-3])
                    run.bold = True
                    run.italic = True
                elif part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    p.add_run(part)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✓ Converted: {os.path.basename(md_file)} → {os.path.basename(docx_file)}")

def main():
    """Convert all markdown files in the root directory to Word format"""
    root_dir = Path('/home/user/application-rationalization-tool')

    # Get all .md files in root directory (not subdirectories)
    md_files = [f for f in root_dir.glob('*.md')]

    if not md_files:
        print("No markdown files found in root directory")
        return

    print(f"Found {len(md_files)} markdown files to convert...\n")

    for md_file in md_files:
        docx_file = md_file.with_suffix('.docx')
        try:
            parse_markdown_to_docx(md_file, docx_file)
        except Exception as e:
            print(f"✗ Error converting {md_file.name}: {e}")

    print(f"\n✓ Conversion complete! Created {len(md_files)} Word documents.")

if __name__ == '__main__':
    main()
