"""
Generate a professionally formatted Word document for the Capital Projects Lifecycle Planner One-Pager
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_one_pager():
    """Create the one-pager Word document"""
    doc = Document()

    # Set document margins (narrower for one-page fit)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('Capital Projects Lifecycle Planner', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('One-Page Executive Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    # What It Is
    heading = doc.add_heading('What It Is', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph('An AI-powered capital project portfolio management platform that transforms infrastructure delivery through automated health monitoring, predictive analytics, and strategic optimization for government agencies and engineering consultants.')
    p.runs[0].font.size = Pt(9)

    # The Problem
    heading = doc.add_heading('The Problem', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('Government agencies waste billions on preventable capital project cost overruns and delays.')
    run.bold = True
    run.font.size = Pt(9)

    problems = [
        'Public works directors manage 50-200+ projects with no unified portfolio health view',
        'Average 25-40% cost overruns on poorly managed infrastructure projects',
        '$100K-1M+ lost per project on preventable delays (ROW, utilities, permitting)',
        '20-40 hours/month wasted manually compiling portfolio status reports',
        'Critical projects slip, grant funding lost due to missed milestones',
        'Project prioritization driven by politics instead of data'
    ]

    for problem in problems:
        p = doc.add_paragraph(problem, style='List Bullet')
        p.runs[0].font.size = Pt(8)

    # The Solution
    heading = doc.add_heading('The Solution', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph('UPLOAD → SCORE → MONITOR → OPTIMIZE')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    # Key Capabilities Table
    p = doc.add_paragraph()
    run = p.add_run('Key Capabilities')
    run.bold = True
    run.font.size = Pt(9)

    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'

    capabilities = [
        ('Feature', 'Benefit'),
        ('Automated Health Scoring', 'Real-time scores (1-100) across 8 dimensions'),
        ('Predictive Analytics', 'Identify problems 3-6 months before they occur'),
        ('ADVANCE Framework', 'Categorize: Advance/Monitor/Re-scope/Defer/Cancel'),
        ('Risk Intelligence', 'Track ROW, utility, permitting, contractor risks'),
        ('Grant Compliance', 'Milestone tracking prevents funding loss'),
        ('AI Assistant', 'Natural language Q&A for portfolio insights'),
        ('What-If Modeling', 'Scenario planning for budget changes'),
        ('Commission-Ready Reports', 'PDF/PowerPoint for executive briefings')
    ]

    for i, (feature, benefit) in enumerate(capabilities):
        row = table.rows[i]
        row.cells[0].text = feature
        row.cells[1].text = benefit

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
                    if i == 0:
                        run.bold = True

    # The Outcome
    heading = doc.add_heading('The Outcome', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    results = [
        '40-60% reduction in preventable cost overruns',
        '80% time savings on portfolio reporting',
        '25-35% improvement in on-time delivery',
        '95%+ grant milestone achievement',
        '15-25% better resource allocation'
    ]

    for result in results:
        p = doc.add_paragraph(result, style='List Bullet')
        p.runs[0].font.size = Pt(8)

    # Sample ROI
    p = doc.add_paragraph()
    run = p.add_run('Sample ROI: ')
    run.bold = True
    run.font.size = Pt(8)
    run = p.add_run('100 Projects | $250M Budget → $6.5M annual value | Investment: $75K-150K | ROI: 44-87x')
    run.font.size = Pt(8)

    # Two column section
    table = doc.add_table(rows=1, cols=2)

    # Who It's For (left column)
    left_cell = table.rows[0].cells[0]
    p = left_cell.paragraphs[0]
    run = p.add_run('Who It\'s For')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = left_cell.add_paragraph()
    run = p.add_run('Government Agencies:')
    run.font.size = Pt(8)
    run.bold = True

    audiences = [
        'County/City Public Works',
        'State DOTs',
        'Transit/Transportation agencies',
        'Water/Wastewater utilities',
        'Ports, airports, special districts'
    ]

    for audience in audiences:
        p = left_cell.add_paragraph(audience, style='List Bullet')
        p.runs[0].font.size = Pt(7)

    p = left_cell.add_paragraph()
    run = p.add_run('Consulting Groups:')
    run.font.size = Pt(8)
    run.bold = True

    consultants = [
        'Civil engineering firms',
        'Program management consultants',
        'Grant compliance consultants',
        'Municipal advisors'
    ]

    for consultant in consultants:
        p = left_cell.add_paragraph(consultant, style='List Bullet')
        p.runs[0].font.size = Pt(7)

    # What Makes It Different (right column)
    right_cell = table.rows[0].cells[1]
    p = right_cell.paragraphs[0]
    run = p.add_run('What Makes It Different')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. Spreadsheets: ')
    run.font.size = Pt(7)
    run.bold = True
    run = p.add_run('Automated, predictive, real-time')
    run.font.size = Pt(7)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. Enterprise PM Tools: ')
    run.font.size = Pt(7)
    run.bold = True
    run = p.add_run('Portfolio intelligence, AI-powered, days to deploy, affordable')
    run.font.size = Pt(7)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. ERP Systems: ')
    run.font.size = Pt(7)
    run.bold = True
    run = p.add_run('Delivery-focused, predictive, actionable')
    run.font.size = Pt(7)

    p = right_cell.add_paragraph()
    run = p.add_run('Unique Strengths:')
    run.font.size = Pt(7)
    run.bold = True

    strengths = [
        'Government infrastructure expertise',
        'Predictive early warning (3-6 mo)',
        'ROW & utility risk intelligence',
        'Grant compliance built-in',
        'Commission/council ready reports',
        'White-label for consultants'
    ]

    for strength in strengths:
        p = right_cell.add_paragraph(strength, style='List Bullet')
        p.runs[0].font.size = Pt(7)

    # Use Cases Table
    heading = doc.add_heading('Proven Use Cases', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light List Accent 1'

    use_cases = [
        ('Scenario', 'Challenge', 'Solution', 'Result'),
        ('Grant Protection', '$15M milestone in 90 days', 'Early warning 4 mo out', '$15M saved'),
        ('Budget Cut', 'Cut $10M from 45 projects', 'Scenario analysis', 'Protected 37 projects'),
        ('New Director', '80 projects, no overview', 'Import in 2 days', 'Found 12 high-risk'),
        ('Consultant Win', '$5M PM contract', 'Demoed AI platform', 'Won vs. larger firms'),
        ('Overrun Prevention', '15% burn rate variance', 'Platform alert', 'Avoided $1.2M overrun')
    ]

    for i, (scenario, challenge, solution, result) in enumerate(use_cases):
        row = table.rows[i]
        row.cells[0].text = scenario
        row.cells[1].text = challenge
        row.cells[2].text = solution
        row.cells[3].text = result

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
                    if i == 0:
                        run.bold = True

    # The Bottom Line
    heading = doc.add_heading('The Bottom Line', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('Transform capital project delivery from reactive chaos to predictive intelligence.')
    run.bold = True
    run.font.size = Pt(9)

    p = doc.add_paragraph('Government agencies can\'t afford to keep losing millions on preventable project failures. This platform provides the visibility, early warning, and strategic optimization needed to deliver infrastructure on-time, on-budget, and with confidence.')
    p.runs[0].font.size = Pt(8)

    # Pricing
    p = doc.add_paragraph()
    run = p.add_run('Pricing: ')
    run.bold = True
    run.font.size = Pt(8)
    run = p.add_run('Government SaaS: $30K-200K/yr (based on project count) | Consulting: $5K-10K/consultant/yr')
    run.font.size = Pt(8)

    # Footer tagline
    p = doc.add_paragraph('"See everything. Predict problems. Optimize outcomes. Transform infrastructure delivery."')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Save the document
    output_path = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\capital_projects\docs\ONE_PAGER.docx'
    doc.save(output_path)
    print(f"SUCCESS: Word document created successfully: {output_path}")

    return output_path

if __name__ == '__main__':
    create_one_pager()
