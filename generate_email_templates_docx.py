"""
Generate a professionally formatted Word document for Email Templates
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_email_templates_docx():
    """Create the email templates Word document"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title = doc.add_heading('Business Outreach Email Templates', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('Platform Capability - Outreach Strategies')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_page_break()

    # Email Option 1
    heading = doc.add_heading('Email Option 1: Direct Enterprise Introduction', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('For Potential Direct Clients')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 102, 102)

    # Subject
    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('Solving Complex Portfolio Management Challenges with Proven AI-Powered Platforms')
    run.font.size = Pt(11)

    # Body
    email_body_1 = """Hi [Name],

I hope this finds you well. I wanted to reach out because I've recently completed something that might be valuable to [Company Name] and your portfolio management challenges.

What I've Built:

Over the past year, I've developed and deployed a universal platform architecture that transforms complex business problems into actionable intelligence. I've successfully implemented this framework across three fundamentally different domains:

1. Application Rationalization - Helping IT organizations optimize their application portfolios (achieved 20-30% cost savings)
2. Capital Projects Management - Enabling government agencies to prevent cost overruns and optimize infrastructure delivery (40-60% reduction in preventable overruns)
3. Contract Oversight - Automating vendor compliance and performance management (10-20% spend optimization)

Why This Matters:

The same architectural pattern that solved these three different problems can be rapidly configured to address virtually any portfolio management challenge - whether it's real estate, fleet management, product portfolios, supplier relationships, or any domain where you need to:
• Manage complex portfolios with multiple attributes
• Make data-driven prioritization decisions
• Track compliance and predict risks
• Optimize spend and resource allocation
• Report to executives with transparency

Proven Results:
• 3-6 month delivery timeline (vs. 12-24 months for traditional custom development)
• 4-87x ROI in first year across implementations
• $150K-500K investment (vs. $1M-5M for traditional approaches)

Next Step:

I'd love to schedule a brief call to discuss whether this framework could address any portfolio management challenges at [Company Name]. I've attached a one-page overview and a comprehensive capability statement for your review.

Would you have 20 minutes in the next week or two for a conversation? I'm happy to work around your schedule.

Best regards,
[Your Name]
[Your Contact Information]

Attachments:
• Platform Capability Statement (Word doc)
• One-Page Overview"""

    for line in email_body_1.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(11)

            # Bold headers
            if line.strip().endswith(':') and len(line.strip()) < 50:
                p.runs[0].bold = True

    doc.add_page_break()

    # Email Option 2
    heading = doc.add_heading('Email Option 2: Consulting Partnership Introduction', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('For Consulting Firms')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 102, 102)

    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('Partnership Opportunity: AI-Powered Platform Architecture for Client Delivery')
    run.font.size = Pt(11)

    email_body_2 = """Hi [Name],

I hope you're doing well. I'm reaching out because I've developed something that could significantly enhance [Firm Name]'s service delivery capabilities and competitive positioning.

The Opportunity:

I've built a proven, replicable platform architecture that rapidly transforms complex portfolio management challenges into AI-powered decision support systems. Over the past year, I've successfully deployed this framework across three different domains with documented ROI of 4-87x.

Why This Could Be Valuable for [Firm Name]:

Rather than delivering recommendations via PowerPoint, your consultants could deliver technology-enabled solutions that:
• Differentiate you from competitors who rely on manual/spreadsheet-based approaches
• Create recurring revenue through SaaS offerings to clients
• Allow you to manage 2-3x more projects per consultant through automation
• Provide measurable, data-driven outcomes that win repeat business

Partnership Models:

I'm open to several collaboration approaches:
1. White-label licensing - You brand it as your own platform
2. Revenue share - I provide the platform, you bring clients, we split revenue
3. Joint venture - Co-develop solutions for your target verticals
4. Technology licensing - You acquire the IP for your exclusive use

Proven Track Record:

Three successful implementations:
• Application Rationalization (IT portfolio optimization)
• Capital Projects Planning (infrastructure delivery)
• Contract Oversight (vendor compliance & performance)

Each delivered in 12-18 weeks with measurable business outcomes. The framework is ready to be applied to your clients' challenges in any domain.

Next Step:

I'd welcome the opportunity to explore how we might work together. I've attached detailed materials, but I think a 30-minute conversation would be most valuable to discuss fit and opportunity.

Are you available for a call in the next week or two?

Best regards,
[Your Name]
[Your Contact Information]

Attachments:
• Platform Capability Statement
• Case Study Summaries"""

    for line in email_body_2.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(11)
            if line.strip().endswith(':') and len(line.strip()) < 50:
                p.runs[0].bold = True

    doc.add_page_break()

    # Email Option 3
    heading = doc.add_heading('Email Option 3: Investor/PE Firm Introduction', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('For Investment Opportunities')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 102, 102)

    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('Portfolio Company Value Creation: Proven Platform Architecture with 4-87x ROI')
    run.font.size = Pt(11)

    email_body_3 = """Hi [Name],

I hope this message finds you well. I'm reaching out because I've developed an intellectual property and delivery capability that could create significant value across [Firm Name]'s portfolio companies.

The Opportunity:

I've built and proven a universal platform architecture that rapidly solves complex portfolio management and operational oversight problems across any industry vertical. Think of it as a meta-capability - a framework that can be quickly configured to deliver data-driven decision support systems for different business domains.

Proven Value Creation:

Three successful implementations with documented outcomes:
1. Application Rationalization → 20-30% IT cost savings
2. Capital Projects Management → 40-60% reduction in cost overruns
3. Contract Oversight → 10-20% spend optimization

Each delivered in 3-6 months with 4-87x first-year ROI.

Portfolio Company Applications:

This framework could rapidly address operational challenges across your holdings:
• Manufacturing companies: Supplier portfolio management, production optimization
• Real estate portfolios: Property portfolio optimization, facilities management
• Healthcare systems: Contract oversight, capital project management
• Retail operations: Product portfolio optimization, vendor management
• Distribution businesses: Fleet management, logistics optimization

Investment Thesis:

• Proven technology (not theoretical - three live implementations)
• Rapid deployment (3-6 months vs. 12-24 for traditional solutions)
• Measurable ROI (documented 4-87x returns)
• Scalable across verticals (same platform, different configurations)
• Recurring revenue model (SaaS + services)

Next Step:

I'd welcome a conversation about:
1. Deploying this capability across your portfolio companies for operational value creation
2. Investment opportunity to scale this as a standalone platform business
3. Partnership to co-develop solutions for specific industries

Would you have 30 minutes for an exploratory discussion? I've attached comprehensive materials for your review.

Best regards,
[Your Name]
[Your Contact Information]

Attachments:
• Platform Capability Statement
• Financial Summary & Business Model"""

    for line in email_body_3.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(11)
            if line.strip().endswith(':') and len(line.strip()) < 50:
                p.runs[0].bold = True

    doc.add_page_break()

    # Email Option 4
    heading = doc.add_heading('Email Option 4: Warm Introduction Request', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('For Mutual Connections')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 102, 102)

    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('Quick Intro Request - AI Platform Architecture for [Target Name/Company]')
    run.font.size = Pt(11)

    email_body_4 = """Hi [Mutual Connection],

Hope you're doing great! I have something I'd love your help with.

I've spent the past year building and proving out a platform architecture that solves complex portfolio management problems across different industries. I've successfully deployed it in three very different domains (IT, infrastructure, procurement) with impressive results - 4-87x ROI, 3-6 month delivery timelines.

The framework is universal - it can be rapidly configured to address portfolio management challenges in virtually any domain (real estate, fleet management, product portfolios, vendor relationships, etc.).

Why I'm reaching out:

I know you have a relationship with [Target Name] at [Target Company], and I believe this capability could be highly relevant to [specific challenge you know they have].

Would you be comfortable making an introduction?

I don't need anything more than an email introduction - something simple like: "Hi [Target], I wanted to connect you with [Your Name] who has built an interesting platform that might be relevant to [challenge/initiative]. Worth a conversation."

I've attached a one-page overview in case you want to preview what I'm working on. Happy to provide more context on a quick call if that's helpful.

Thanks for considering! Let me know if you need anything else.

Best,
[Your Name]

P.S. - If you know anyone else in your network dealing with complex portfolio management challenges, I'd appreciate any introductions!"""

    for line in email_body_4.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(11)
            if line.strip().endswith(':') and len(line.strip()) < 50:
                p.runs[0].bold = True

    doc.add_page_break()

    # Email Option 5
    heading = doc.add_heading('Email Option 5: LinkedIn Message Version', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('Character Limit Friendly')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 102, 102)

    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('Proven Platform for Portfolio Management Challenges')
    run.font.size = Pt(11)

    email_body_5 = """Hi [Name],

I came across your profile and wanted to reach out based on your work in [their domain/industry].

I've built a universal platform architecture that's proven across 3 different domains (IT portfolio, infrastructure projects, contract management) - each achieving 4-87x ROI in 3-6 months.

The framework rapidly solves portfolio management challenges across any industry:
✓ Data-driven decision support
✓ Predictive analytics & risk detection
✓ Compliance automation
✓ Strategic optimization

Given your experience with [specific challenge], I thought this might be relevant.

Open to a brief conversation if you're interested. I can share detailed materials.

Best regards,
[Your Name]"""

    for line in email_body_5.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # Email Option 6
    heading = doc.add_heading('Email Option 6: Mass Outreach / Newsletter', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('For Broader Networking')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 102, 102)

    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('Announcing: Universal Platform Architecture for Business Portfolio Challenges')
    run.font.size = Pt(11)

    email_body_6 = """Hi everyone,

I wanted to share something I've been working on that might be valuable to many of you.

What I've Built:

A proven platform architecture that rapidly transforms complex portfolio management challenges into AI-powered decision support systems. Think of it as a universal framework that can be configured to solve different business problems across industries.

Proven Track Record:

Three successful implementations in the past year:
• Application Rationalization (IT portfolio optimization) → 20-30% cost savings
• Capital Projects Management (infrastructure delivery) → 40-60% reduction in cost overruns
• Contract Oversight (vendor compliance) → 10-20% spend optimization

All delivered in 3-6 months with 4-87x first-year ROI.

Universal Applications:

The same architecture applies to virtually any portfolio management challenge:
Real estate portfolios • Fleet & equipment management • Product portfolio optimization • Supplier/vendor relationships • R&D project portfolios • Grant & fund management • Risk management • Investment portfolios

Who This Might Help:

• Enterprise leaders struggling with manual processes and lack of visibility
• Consulting firms looking to enhance service delivery with technology
• PE/VC firms seeking operational value creation across portfolio companies
• Technology companies wanting to expand their product offerings

Next Steps:

If this resonates with a challenge you're facing (or know someone who is), I'd love to have a conversation. I've created comprehensive materials that detail the framework, case studies, and engagement models.

Reply to this email or schedule time on my calendar: [Calendar Link]

Attachments:
• One-Page Overview
• Platform Capability Statement

Best regards,
[Your Name]
[Your Title/Company]
[Contact Information]"""

    for line in email_body_6.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(11)
            if line.strip().endswith(':') and len(line.strip()) < 50:
                p.runs[0].bold = True

    doc.add_page_break()

    # Follow-up Email
    heading = doc.add_heading('Follow-Up Email Template', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('7-10 Days After Initial Outreach')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 102, 102)

    p = doc.add_paragraph()
    run = p.add_run('Subject: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('Re: [Original Subject] - Following Up')
    run.font.size = Pt(11)

    followup = """Hi [Name],

I wanted to follow up on my email from [date] about the platform architecture I've developed for portfolio management challenges.

I know your inbox is likely full, so I'll keep this brief:

Quick Recap:
• Universal platform proven across 3 different domains
• 4-87x ROI in 3-6 months
• Applies to virtually any portfolio management challenge

Thought I'd Add:

I recently came across [recent news/article about their company or industry] and it reinforced why this capability might be timely for [Company Name]. [Specific insight about how your platform addresses a challenge mentioned in the news].

No Pressure:

If the timing isn't right or it's not a fit, totally understand. But if you'd like to explore further, I'm happy to:
• Send additional materials
• Schedule a brief 15-minute call
• Connect you with references from my implementations

Just let me know what works best.

Best regards,
[Your Name]"""

    for line in followup.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(11)
            if line.strip().endswith(':') and len(line.strip()) < 50:
                p.runs[0].bold = True

    doc.add_page_break()

    # Best Practices Section
    heading = doc.add_heading('Email Best Practices & Tips', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    subheading = doc.add_heading('Personalization is Key', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    tips_1 = [
        'Replace all [bracketed] placeholders with specific details',
        'Research the recipient - reference their recent work, company initiatives, or industry challenges',
        'Adjust tone based on your relationship (formal for cold outreach, casual for close contacts)'
    ]

    for tip in tips_1:
        p = doc.add_paragraph(tip, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    subheading = doc.add_heading('Timing Matters', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    tips_2 = [
        'Best send times: Tuesday-Thursday, 10am-2pm local time',
        'Avoid: Monday mornings, Friday afternoons, holidays',
        'Follow up: 7-10 days if no response, then again at 2-3 weeks'
    ]

    for tip in tips_2:
        p = doc.add_paragraph(tip, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    subheading = doc.add_heading('Subject Line Tips', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    tips_3 = [
        'Keep under 50 characters',
        'Be specific (avoid "Opportunity" or "Quick Question")',
        'Create urgency or curiosity when appropriate',
        'Consider A/B testing different approaches'
    ]

    for tip in tips_3:
        p = doc.add_paragraph(tip, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    subheading = doc.add_heading('Call-to-Action Best Practices', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    tips_4 = [
        'Be specific: "Are you available Tuesday or Wednesday next week?"',
        'Make it easy: Include calendar link',
        'Offer flexibility: "Happy to work around your schedule"',
        'Lower the barrier: "Just 15-20 minutes" vs. "an hour"'
    ]

    for tip in tips_4:
        p = doc.add_paragraph(tip, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # Response Handling
    heading = doc.add_heading('Response Handling Scripts', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Response 1
    subheading = doc.add_heading('If They Say "Send Me More Information"', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    response_1 = """Absolutely! I've attached our comprehensive Platform Capability Statement that details the framework, case studies, and engagement models. I've also included a one-page executive summary for quick reference.

After you've had a chance to review, would you like to schedule a brief call to discuss how this might apply to [specific challenge]? I'm available [specific times]."""

    p = doc.add_paragraph(response_1)
    p.runs[0].font.size = Pt(11)

    # Response 2
    subheading = doc.add_heading('If They Say "Not Right Now"', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    response_2 = """Completely understand - timing is everything. Would it make sense to reconnect in [3/6] months? Or if you know anyone in your network dealing with [portfolio management challenges], I'd appreciate any introductions you're comfortable making."""

    p = doc.add_paragraph(response_2)
    p.runs[0].font.size = Pt(11)

    # Response 3
    subheading = doc.add_heading('If They Say "Tell Me More"', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    response_3 = """Great! The short version: I've built a universal platform that rapidly solves portfolio management problems across any domain. Proven in 3 different industries with 4-87x ROI.

But rather than me explaining over email, would a 15-minute call work better? I can share screen and show you exactly how it works. Are you available [specific times]?"""

    p = doc.add_paragraph(response_3)
    p.runs[0].font.size = Pt(11)

    # Response 4
    subheading = doc.add_heading('If They Ask About Pricing', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    response_4 = """Investment typically ranges from $150K-500K depending on portfolio size and complexity - significantly less than traditional custom development ($1M-5M) or enterprise software implementations.

The ROI has been 4-87x in first year across our implementations. But the specific investment would depend on your unique needs. Worth a conversation to explore fit and scope?"""

    p = doc.add_paragraph(response_4)
    p.runs[0].font.size = Pt(11)

    # Save the document
    output_path = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\EMAIL_TEMPLATE_BUSINESS_ASSOCIATES.docx'
    doc.save(output_path)
    print(f"SUCCESS: Word document created successfully: {output_path}")

    return output_path

if __name__ == '__main__':
    create_email_templates_docx()
