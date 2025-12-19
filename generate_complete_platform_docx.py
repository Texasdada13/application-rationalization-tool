"""
Generate a complete, properly formatted Word document from Platform Capability Statement markdown
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def process_markdown_to_docx(md_file_path, output_path):
    """Process markdown file and create Word document with proper formatting"""

    # Read the markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set document margins - standard for readability
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Split content into lines
    lines = content.split('\n')

    in_table = False
    table = None
    table_headers = []
    table_row_idx = 0
    skip_next = False

    i = 0
    while i < len(lines):
        if skip_next:
            skip_next = False
            i += 1
            continue

        line = lines[i].rstrip()

        # Skip horizontal rules
        if line.strip() == '---':
            if i > 0 and i < len(lines) - 1:
                doc.add_paragraph()  # Add spacing
            i += 1
            continue

        # Handle tables
        if '|' in line and not in_table:
            # Check if this is actually a table (has separator line after header)
            if i + 1 < len(lines) and re.match(r'^[\s\|:\-]+$', lines[i + 1]):
                parts = [p.strip() for p in line.split('|') if p.strip()]
                if parts:
                    in_table = True
                    table_headers = parts

                    # Count table rows
                    row_count = 1  # header
                    for j in range(i + 2, len(lines)):
                        if '|' in lines[j] and not lines[j].strip().startswith('#') and not re.match(r'^[\s\|:\-]+$', lines[j]):
                            row_count += 1
                        elif lines[j].strip() and not lines[j].strip().startswith('|'):
                            break

                    # Create table
                    table = doc.add_table(rows=row_count, cols=len(table_headers))
                    table.style = 'Light Grid Accent 1'

                    # Add headers
                    for col_idx, header in enumerate(table_headers):
                        cell = table.rows[0].cells[col_idx]
                        # Remove ** formatting from headers
                        header_text = re.sub(r'\*\*(.*?)\*\*', r'\1', header)
                        cell.text = header_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)

                    table_row_idx = 1
                    i += 1  # Skip to separator line
                    continue

        if in_table and '|' in line:
            # Skip separator line
            if re.match(r'^[\s\|:\-]+$', line):
                i += 1
                continue

            # Add table row
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if parts and table and table_row_idx < len(table.rows):
                for col_idx, cell_text in enumerate(parts):
                    if col_idx < len(table.rows[table_row_idx].cells):
                        cell = table.rows[table_row_idx].cells[col_idx]
                        # Remove markdown formatting
                        cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                table_row_idx += 1
                i += 1
                continue
            else:
                # End of table
                in_table = False
                table = None
                doc.add_paragraph()  # Add spacing after table

        # Not in table or table ended
        if in_table and '|' not in line:
            in_table = False
            table = None
            doc.add_paragraph()

        # H1 - Main sections
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if 'Platform Capability Statement' in text:
                # Title page
                doc.add_paragraph()  # Top spacing
                title = doc.add_heading(text.split(':')[0], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.size = Pt(28)
                title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

                if ':' in text:
                    subtitle = doc.add_paragraph(text.split(':', 1)[1].strip())
                    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    subtitle.runs[0].font.size = Pt(18)
                    subtitle.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                    subtitle.runs[0].bold = True

                doc.add_page_break()
            else:
                # Regular H1 - Add page break before major sections for better flow
                if i > 10:  # Not at the beginning
                    doc.add_page_break()
                heading = doc.add_heading(text, 1)
                heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                heading.runs[0].font.size = Pt(16)

        # H2 - Subsections
        elif line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, 2)
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            heading.runs[0].font.size = Pt(14)

        # H3 - Sub-subsections
        elif line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, 3)
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            heading.runs[0].font.size = Pt(12)

        # H4 - Minor headings
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)

        # Bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]

            p = doc.add_paragraph(style='List Bullet')

            # Handle inline bold formatting
            if '**' in text:
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(11)
                    elif part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
            else:
                p.add_run(text)
                p.runs[0].font.size = Pt(11)

        # Numbered lists
        elif re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())

            p = doc.add_paragraph(style='List Number')

            # Handle inline bold
            if '**' in text:
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(11)
                    elif part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
            else:
                p.add_run(text)
                p.runs[0].font.size = Pt(11)

        # Code blocks - render as center-aligned, bold text
        elif line.strip().startswith('```') and not line.strip().endswith('```'):
            # Multi-line code block start
            i += 1
            code_content = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_content.append(lines[i].strip())
                i += 1
            if code_content:
                p = doc.add_paragraph('\n'.join(code_content))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(12)
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

        # Single line code
        elif line.strip().startswith('```') and line.strip().endswith('```'):
            text = line.strip()[3:-3]
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(12)
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

        # Bold standalone paragraph
        elif line.strip().startswith('**') and line.strip().endswith('**') and ':' not in line:
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)

        # Regular paragraphs
        elif line.strip() and not line.startswith('#') and not in_table:
            # Check if it's a label (ends with colon and is short)
            if line.strip().endswith(':') and len(line.strip()) < 80:
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.bold = True
                run.font.size = Pt(11)
            else:
                # Regular paragraph with potential inline formatting
                p = doc.add_paragraph()

                if '**' in line:
                    parts = re.split(r'(\*\*.*?\*\*)', line.strip())
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                            run.font.size = Pt(11)
                        elif part:
                            run = p.add_run(part)
                            run.font.size = Pt(11)
                else:
                    run = p.add_run(line.strip())
                    run.font.size = Pt(11)

        # Empty lines - add spacing
        elif not line.strip() and not in_table:
            # Only add paragraph breaks in appropriate contexts
            if i > 0 and i < len(lines) - 1:
                if not lines[i-1].startswith('#') and not lines[i+1].startswith('#'):
                    pass  # Word handles spacing

        i += 1

    # Save document
    doc.save(output_path)
    print(f"SUCCESS: Complete Word document created: {output_path}")
    print(f"Document contains:")
    print(f"  - {len(doc.paragraphs)} paragraphs")
    print(f"  - {len(doc.tables)} tables")
    print(f"  - Processed {len(lines)} lines from markdown")

if __name__ == '__main__':
    md_file = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\PLATFORM_CAPABILITY_STATEMENT.md'
    output_file = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\PLATFORM_CAPABILITY_STATEMENT.docx'

    process_markdown_to_docx(md_file, output_file)
