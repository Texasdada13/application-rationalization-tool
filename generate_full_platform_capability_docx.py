"""
Generate a complete Word document from the Platform Capability Statement markdown file
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import re

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def process_markdown_to_docx(md_file_path, output_path):
    """Process markdown file and create Word document"""

    # Read the markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Split content into lines
    lines = content.split('\n')

    in_table = False
    table = None
    table_headers = []

    for i, line in enumerate(lines):
        line = line.rstrip()

        # Skip horizontal rules
        if line.strip() == '---':
            if i > 0 and i < len(lines) - 1:  # Not at start or end
                doc.add_paragraph()  # Just add spacing
            continue

        # Handle tables
        if '|' in line and not in_table:
            # Start of table
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if parts:
                in_table = True
                table_headers = parts
                # Look ahead to count rows
                row_count = 1  # header
                for j in range(i+2, len(lines)):  # Skip separator line
                    if '|' in lines[j] and not lines[j].strip().startswith('#'):
                        row_count += 1
                    else:
                        break

                # Create table
                table = doc.add_table(rows=row_count, cols=len(table_headers))
                table.style = 'Light Grid Accent 1'

                # Add headers
                for col_idx, header in enumerate(table_headers):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)

                table_row_idx = 1
            continue

        elif '|' in line and in_table:
            # Check if it's the separator line
            if re.match(r'^[\s\|:\-]+$', line):
                continue

            # Add table row
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if parts and table and table_row_idx < len(table.rows):
                for col_idx, cell_text in enumerate(parts):
                    if col_idx < len(table.rows[table_row_idx].cells):
                        cell = table.rows[table_row_idx].cells[col_idx]
                        # Remove markdown formatting
                        cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                table_row_idx += 1
            continue

        else:
            if in_table:
                in_table = False
                table = None
                table_row_idx = 0

        # H1 - Main sections
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if 'Platform Capability Statement' in text:
                # Title page
                title = doc.add_heading(text.split(':')[0], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.size = Pt(24)
                title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

                if ':' in text:
                    subtitle = doc.add_paragraph(text.split(':', 1)[1].strip())
                    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    subtitle.runs[0].font.size = Pt(16)
                    subtitle.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                    subtitle.runs[0].bold = True

                add_page_break(doc)
            else:
                # Regular H1
                heading = doc.add_heading(text, 1)
                heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                heading.runs[0].font.size = Pt(14)

        # H2 - Subsections
        elif line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, 2)
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            heading.runs[0].font.size = Pt(12)

        # H3 - Sub-subsections
        elif line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, 3)
            heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            heading.runs[0].font.size = Pt(11)

        # H4 - Minor headings
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)

        # Bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', lambda m: m.group(1), text)

            p = doc.add_paragraph(text, style='List Bullet')
            p.runs[0].font.size = Pt(10)

            # Apply bold to parts that were **bold**
            if '**' in line:
                p.runs[0].bold = False
                parts = re.split(r'(\*\*.*?\*\*)', line.strip()[2:])
                p.clear()
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(10)

        # Numbered lists
        elif re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            # Remove markdown formatting for text
            text_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)

            p = doc.add_paragraph(text_clean, style='List Number')
            p.runs[0].font.size = Pt(10)

            # Apply bold to **text**
            if '**' in text:
                p.runs[0].bold = False
                p.clear()
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(10)

        # Code blocks
        elif line.strip().startswith('```'):
            continue  # Skip code block markers

        # Bold text paragraph (starts with **)
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)

        # Regular paragraphs
        elif line.strip() and not line.startswith('#'):
            # Check if this is part of a definition list (ends with :)
            if line.strip().endswith(':') and len(line.strip()) < 100:
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.bold = True
                run.font.size = Pt(10)
            else:
                p = doc.add_paragraph(line.strip())
                p.runs[0].font.size = Pt(10)

                # Handle inline bold
                if '**' in line:
                    p.clear()
                    parts = re.split(r'(\*\*.*?\*\*)', line.strip())
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                            run.font.size = Pt(10)
                        else:
                            run = p.add_run(part)
                            run.font.size = Pt(10)

        # Empty lines
        else:
            # Add a small paragraph for spacing, but not after headings
            if i > 0 and not lines[i-1].startswith('#'):
                pass  # Word handles spacing automatically

    # Save document
    doc.save(output_path)
    print(f"SUCCESS: Complete Word document created: {output_path}")
    print(f"Processed {len(lines)} lines from markdown file")

if __name__ == '__main__':
    md_file = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\PLATFORM_CAPABILITY_STATEMENT.md'
    output_file = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\PLATFORM_CAPABILITY_STATEMENT.docx'

    process_markdown_to_docx(md_file, output_file)
