"""
Generate a professionally formatted Word document for the Contract Oversight System One-Pager
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_one_pager():
    """Create the one-pager Word document"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('Contract Oversight System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('One-Page Executive Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    # What It Is
    heading = doc.add_heading('What It Is', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph('An intelligent contract lifecycle management platform that automates compliance monitoring, tracks vendor performance, and provides predictive risk intelligence to maximize contract value and prevent failures across government and enterprise contract portfolios.')
    p.runs[0].font.size = Pt(9)

    # The Problem
    heading = doc.add_heading('The Problem', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('Organizations waste billions on contract mismanagement, compliance violations, and vendor underperformance.')
    run.bold = True
    run.font.size = Pt(9)

    problems = [
        'Manage 200-5,000+ contracts with no unified compliance/performance view',
        '$50K-500K+ lost per compliance violation',
        '40-60% of contract manager time wasted on manual admin',
        'Missed renewals cost $100K-1M+ in unfavorable terms',
        'Vendor performance issues discovered after failures',
        'Contract data scattered - no portfolio analytics'
    ]

    for problem in problems:
        p = doc.add_paragraph(problem, style='List Bullet')
        p.runs[0].font.size = Pt(8)

    # The Solution
    heading = doc.add_heading('The Solution', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph('IMPORT → MONITOR → ANALYZE → OPTIMIZE')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    # Key Capabilities Table
    p = doc.add_paragraph()
    run = p.add_run('Key Capabilities')
    run.bold = True
    run.font.size = Pt(9)

    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'

    capabilities = [
        ('Feature', 'Benefit'),
        ('Automated Compliance', 'Track insurance, licenses, deliverables with alerts'),
        ('Vendor Performance', 'Systematic scoring, benchmarking, historical tracking'),
        ('Risk Intelligence', 'AI predicts failures 3-6 months ahead'),
        ('Spend Analytics', 'Portfolio visibility, consolidation opportunities'),
        ('Audit-Ready Reports', 'One-click compliance reports for regulations'),
        ('Vendor Portal', 'Self-service document submission'),
        ('AI Assistant', 'Natural language Q&A for portfolio insights'),
        ('Workflow Automation', 'Approval routing, renewals, escalations')
    ]

    for i, (feature, benefit) in enumerate(capabilities):
        row = table.rows[i]
        row.cells[0].text = feature
        row.cells[1].text = benefit

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
                    if i == 0:
                        run.bold = True

    # The Outcome
    heading = doc.add_heading('The Outcome', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    results = [
        '10-20% contract spend savings through optimization',
        '80-95% reduction in audit findings',
        '60-80% time savings on contract administration',
        '90%+ on-time renewals eliminating gaps',
        '25-40% improvement in vendor performance',
        '90% faster audit preparation'
    ]

    for result in results:
        p = doc.add_paragraph(result, style='List Bullet')
        p.runs[0].font.size = Pt(8)

    # Sample ROI
    p = doc.add_paragraph()
    run = p.add_run('Sample ROI: ')
    run.bold = True
    run.font.size = Pt(8)
    run = p.add_run('500 Contracts | $150M Spend → $9.8M annual value | Investment: $150K-300K | ROI: 33-66x')
    run.font.size = Pt(8)

    # Two column section
    table = doc.add_table(rows=1, cols=2)

    # Who It's For (left column)
    left_cell = table.rows[0].cells[0]
    p = left_cell.paragraphs[0]
    run = p.add_run('Who It\'s For')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = left_cell.add_paragraph()
    run = p.add_run('Government:')
    run.font.size = Pt(8)
    run.bold = True

    gov = [
        'County/City procurement',
        'State agencies',
        'School districts',
        'Higher education',
        'Utilities, transit agencies'
    ]

    for item in gov:
        p = left_cell.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(7)

    p = left_cell.add_paragraph()
    run = p.add_run('Private Sector:')
    run.font.size = Pt(8)
    run.bold = True

    private = [
        'Construction companies',
        'Healthcare systems',
        'Manufacturers',
        'Financial institutions',
        'Real estate/property mgmt'
    ]

    for item in private:
        p = left_cell.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(7)

    # What Makes It Different (right column)
    right_cell = table.rows[0].cells[1]
    p = right_cell.paragraphs[0]
    run = p.add_run('What Makes It Different')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. Spreadsheets: ')
    run.font.size = Pt(7)
    run.bold = True
    run = p.add_run('Automated, proactive, analytics, vendor portals')
    run.font.size = Pt(7)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. Generic CLM: ')
    run.font.size = Pt(7)
    run.bold = True
    run = p.add_run('Compliance focus, performance depth, risk intelligence')
    run.font.size = Pt(7)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. ERP Systems: ')
    run.font.size = Pt(7)
    run.bold = True
    run = p.add_run('Oversight depth, not just transactions')
    run.font.size = Pt(7)

    p = right_cell.add_paragraph()
    run = p.add_run('Unique Strengths:')
    run.font.size = Pt(7)
    run.bold = True

    strengths = [
        'Compliance-first design',
        'Vendor performance engine',
        'Predictive risk intelligence',
        'Government procurement expertise',
        'Insurance & license tracking',
        'Audit readiness',
        'White-label capable'
    ]

    for strength in strengths:
        p = right_cell.add_paragraph(strength, style='List Bullet')
        p.runs[0].font.size = Pt(7)

    # Use Cases Table
    heading = doc.add_heading('Proven Use Cases', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light List Accent 1'

    use_cases = [
        ('Scenario', 'Challenge', 'Solution', 'Result'),
        ('Insurance Gap', 'Lapse, $2M lawsuit', 'Auto alerts 30/60/90 days', '$2M avoided'),
        ('Audit', 'Federal audit, 3 weeks prep', 'One-click reports', 'Zero findings, saved 300 hrs'),
        ('Performance Defense', 'Procurement protest', '3-year performance data', 'Protest dismissed'),
        ('Spend Optimization', 'CFO demands $5M savings', 'Analytics identify opportunities', '$6.5M found'),
        ('Consulting Scale', 'Can\'t grow without staff', 'White-label, 70% automation', '5 to 12 clients, 40% profit'),
        ('Grant Compliance', '$100M grants, 20 projects', 'Automated tracking', '100% on-time, zero violations')
    ]

    for i, (scenario, challenge, solution, result) in enumerate(use_cases):
        row = table.rows[i]
        row.cells[0].text = scenario
        row.cells[1].text = challenge
        row.cells[2].text = solution
        row.cells[3].text = result

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6)
                    if i == 0:
                        run.bold = True

    # The Bottom Line
    heading = doc.add_heading('The Bottom Line', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    heading.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('Transform contract oversight from reactive chaos to predictive intelligence.')
    run.bold = True
    run.font.size = Pt(9)

    p = doc.add_paragraph('Organizations can\'t afford contract mismanagement - the risks are too high. This platform provides automated compliance monitoring, performance intelligence, and strategic insights to maximize contract value while minimizing risk.')
    p.runs[0].font.size = Pt(8)

    # Pricing
    p = doc.add_paragraph()
    run = p.add_run('Pricing: ')
    run.bold = True
    run.font.size = Pt(8)
    run = p.add_run('Government/Private SaaS: $50K-300K/yr | Consulting: $3K-5K/consultant/yr | White-label available')
    run.font.size = Pt(8)

    # Footer tagline
    p = doc.add_paragraph('"Ensure compliance. Optimize performance. Maximize value. Transform contract oversight."')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Save the document
    output_path = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\contract_oversight_system\docs\ONE_PAGER.docx'
    doc.save(output_path)
    print(f"SUCCESS: Word document created successfully: {output_path}")

    return output_path

if __name__ == '__main__':
    create_one_pager()
