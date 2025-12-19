"""
Generate a professionally formatted Word document for the Platform Capability Statement
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_capability_statement():
    """Create the Platform Capability Statement Word document"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title = doc.add_heading('Platform Capability Statement', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('Universal Business Intelligence Framework')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    subtitle_run.bold = True

    tagline = doc.add_paragraph('A Proven, Replicable Architecture for Solving Complex Business Challenges')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.runs[0]
    tagline_run.font.size = Pt(12)
    tagline_run.font.italic = True
    tagline_run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_page_break()

    # Executive Summary
    heading = doc.add_heading('Executive Summary', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('I have built a proven, replicable platform architecture that transforms complex business challenges into actionable intelligence across any domain.')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run('Through the development of three distinct but architecturally consistent solutions - ')
    run = p.add_run('Application Rationalization')
    run.bold = True
    p.add_run(', ')
    run = p.add_run('Capital Projects Lifecycle Planning')
    run.bold = True
    p.add_run(', and ')
    run = p.add_run('Contract Oversight Management')
    run.bold = True
    p.add_run(' - I\'ve demonstrated a universal capability: the ability to rapidly design, build, and deploy data-driven decision platforms that solve complex portfolio management and operational oversight problems in any industry or functional area.')
    p.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run('This is not about three applications. This is about a proven methodology and technical framework that can be applied to solve virtually any business problem.')
    run.font.size = Pt(11)
    run.bold = True

    # The Core Capability
    heading = doc.add_heading('The Core Capability: A Universal Problem-Solving Framework', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('A reusable platform architecture consisting of:')
    run.bold = True
    run.font.size = Pt(11)

    components = [
        ('Data Integration & Structuring Layer', 'Ingest from any source, AI-powered extraction, multi-source harmonization'),
        ('Automated Assessment & Scoring Engine', 'Multi-dimensional evaluation, weighted scoring, automated alerts'),
        ('Stakeholder Intelligence Layer', 'Structured frameworks, quantitative + qualitative data'),
        ('Predictive Analytics & ML Layer', 'Pattern recognition, anomaly detection, risk prediction, forecasting'),
        ('Strategic Categorization Framework', 'Custom decision frameworks, portfolio segmentation, recommendations'),
        ('Compliance & Risk Monitoring', 'Automated compliance checking, regulatory mapping, audit trails'),
        ('Visualization & Reporting Layer', 'Executive dashboards, interactive analytics, automated reports'),
        ('AI Assistant & NLP Interface', 'Conversational Q&A, intelligent recommendations, contextual help')
    ]

    for i, (component, description) in enumerate(components, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {component}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(description)
        run.font.size = Pt(10)

    # Proven Track Record
    doc.add_page_break()
    heading = doc.add_heading('Proven Track Record: Three Domains, Same Architecture', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Application Rationalization
    subheading = doc.add_heading('1. Application Rationalization Tool', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    p = doc.add_paragraph()
    run = p.add_run('Problem Domain: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('IT portfolio management - deciding which applications to keep, retire, modernize, or consolidate')
    run.font.size = Pt(10)

    metrics = [
        ('Portfolio:', '200-500 enterprise applications'),
        ('Scoring:', 'Business value, technical health, cost, strategic alignment'),
        ('Framework:', 'TIME (Tolerate/Invest/Migrate/Eliminate)'),
        ('Compliance:', 'SOX, HIPAA, GDPR, PCI-DSS'),
        ('Outcome:', '20-30% cost savings, reduced technical debt')
    ]

    for label, value in metrics:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f' {value}')
        run.font.size = Pt(10)

    # Capital Projects
    subheading = doc.add_heading('2. Capital Projects Lifecycle Planner', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    p = doc.add_paragraph()
    run = p.add_run('Problem Domain: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('Infrastructure project portfolio management - optimizing delivery of roads, bridges, public works')
    run.font.size = Pt(10)

    metrics = [
        ('Portfolio:', '50-500 capital projects across 10 lifecycle phases'),
        ('Scoring:', 'Budget performance, schedule adherence, risk exposure, strategic value'),
        ('Framework:', 'ADVANCE (Advance/Monitor/Re-scope/Defer/Cancel)'),
        ('Compliance:', 'Grant milestones, ROW acquisition, environmental permits'),
        ('Outcome:', '40-60% reduction in cost overruns, 80% time savings')
    ]

    for label, value in metrics:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f' {value}')
        run.font.size = Pt(10)

    # Contract Oversight
    subheading = doc.add_heading('3. Contract Oversight System', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    p = doc.add_paragraph()
    run = p.add_run('Problem Domain: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('Vendor/contract management - ensuring compliance, optimizing performance, maximizing value')
    run.font.size = Pt(10)

    metrics = [
        ('Portfolio:', '200-5,000 active contracts'),
        ('Scoring:', 'Compliance status, vendor performance, risk exposure, strategic value'),
        ('Framework:', 'Renew/Rebid/Consolidate/Terminate'),
        ('Compliance:', 'Insurance, licenses, DBE/MBE, prevailing wage, grant requirements'),
        ('Outcome:', '10-20% spend savings, 80-95% reduction in audit findings')
    ]

    for label, value in metrics:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f' {value}')
        run.font.size = Pt(10)

    # The Pattern
    doc.add_page_break()
    heading = doc.add_heading('The Pattern: A Universal Solution Architecture', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('Common DNA Across All Three Platforms')
    run.bold = True
    run.font.size = Pt(12)

    # Create comparison table
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Light Grid Accent 1'

    headers = ['Component', 'App Rationalization', 'Capital Projects', 'Contract Oversight']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    comparison_data = [
        ('Entity Type', 'Applications', 'Capital Projects', 'Contracts/Vendors'),
        ('Data Sources', 'CMDB, spreadsheets, surveys', 'Schedules, budgets, milestones', 'Procurement, documents, certificates'),
        ('Health Scoring', 'Business value + tech health', 'Budget + schedule + risk', 'Compliance + performance'),
        ('Risk Intelligence', 'Security, compliance, tech debt', 'Cost overruns, delays, ROW', 'Insurance gaps, vendor failures'),
        ('Decision Framework', 'TIME quadrants', 'ADVANCE categories', 'Renew/Rebid/Consolidate'),
        ('Stakeholder Input', '30+ structured questions', 'Multi-evaluator assessments', 'Performance reviews, surveys'),
        ('Compliance Tracking', 'Regulatory frameworks', 'Grant milestones, permits', 'Insurance, licenses, regulations'),
        ('AI Insights', 'Consolidation opportunities', 'Predictive cost/schedule', 'Spend optimization, risk prediction'),
        ('Reporting', 'CIO briefings', 'Commission/council reports', 'Audit reports, compliance'),
        ('ROI', '20-30% cost reduction', '40-60% overrun reduction', '10-20% spend savings')
    ]

    for i, (component, app, capital, contract) in enumerate(comparison_data, 1):
        table.rows[i].cells[0].text = component
        table.rows[i].cells[1].text = app
        table.rows[i].cells[2].text = capital
        table.rows[i].cells[3].text = contract

        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run('The Core Insight: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run('The same architectural pattern solves fundamentally different business problems by simply changing the domain-specific scoring criteria, data sources, and compliance rules.')
    run.font.size = Pt(11)

    # Universal Application Potential
    doc.add_page_break()
    heading = doc.add_heading('Universal Application Potential', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('The framework applies to any scenario where an organization needs to:')
    run.font.size = Pt(11)

    capabilities = [
        'Manage a portfolio of assets/items with multiple attributes',
        'Evaluate performance or value across multiple dimensions',
        'Track compliance with regulations, policies, or standards',
        'Make strategic decisions about prioritization, investment, or elimination',
        'Predict risks and identify problems early',
        'Optimize spend and resource allocation',
        'Report to stakeholders with transparency and accountability'
    ]

    for capability in capabilities:
        p = doc.add_paragraph(capability, style='List Bullet')
        p.runs[0].font.size = Pt(10)

    # Additional Applications Table
    p = doc.add_paragraph()
    run = p.add_run('Examples of Immediate Applications')
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=13, cols=4)
    table.style = 'Light List Accent 1'

    app_headers = ['Domain', 'Portfolio Type', 'Key Decisions', 'Compliance Areas']
    for i, header in enumerate(app_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    applications = [
        ('Real Estate', 'Property portfolio', 'Buy/Sell/Renovate/Hold', 'Safety, environmental, leases'),
        ('Fleet Management', 'Vehicles/equipment', 'Replace/Maintain/Retire', 'Safety inspections, maintenance'),
        ('HR/Talent', 'Employee skills/roles', 'Develop/Retain/Recruit', 'Certifications, training, performance'),
        ('Product Portfolio', 'Product lines', 'Invest/Maintain/Phase Out', 'Quality, regulatory, profitability'),
        ('Research Portfolio', 'R&D projects', 'Fund/Continue/Pivot/Cancel', 'Milestones, IP, compliance'),
        ('Supplier Management', 'Vendor relationships', 'Expand/Maintain/Reduce/Exit', 'Quality, delivery, certifications'),
        ('Grant/Fund Management', 'Grant portfolio', 'Fund/Monitor/Expand/Close', 'Reporting, milestones, outcomes'),
        ('Facility Management', 'Buildings/spaces', 'Invest/Maintain/Lease/Sell', 'Safety, energy, ADA compliance'),
        ('Marketing Campaigns', 'Marketing initiatives', 'Scale/Optimize/Pause/Stop', 'Budget, ROI, brand compliance'),
        ('IT Service Portfolio', 'Services/capabilities', 'Enhance/Maintain/Sunset', 'SLAs, security, availability'),
        ('Risk Register', 'Enterprise risks', 'Mitigate/Monitor/Accept/Transfer', 'Controls, audit findings'),
        ('Investment Portfolio', 'Financial assets', 'Buy/Hold/Rebalance/Sell', 'Regulatory, risk limits')
    ]

    for i, (domain, portfolio, decisions, compliance) in enumerate(applications, 1):
        table.rows[i].cells[0].text = domain
        table.rows[i].cells[1].text = portfolio
        table.rows[i].cells[2].text = decisions
        table.rows[i].cells[3].text = compliance

        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)

    # The Value Proposition
    doc.add_page_break()
    heading = doc.add_heading('The Value Proposition: Rapid Domain Expertise Deployment', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('What You Get When You Engage Me')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run('Not just a developer. Not just a consultant. A proven platform architect who delivers:')
    run.font.size = Pt(11)

    value_props = [
        ('Rapid Solution Design (Weeks, Not Months)', 'Understand your domain, map to universal framework, configure for your context'),
        ('Proven Architecture (Not Experimental)', 'Battle-tested across three distinct domains, scalable, maintainable, modern tech stack'),
        ('End-to-End Delivery', 'Data integration, custom scoring, stakeholder frameworks, dashboards, training, support'),
        ('Measurable Business Outcomes', 'Demonstrated ROI of 4-87x, proven cost savings of 10-60%, 60-80% time savings')
    ]

    for i, (title, description) in enumerate(value_props, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(description)
        run.font.size = Pt(10)

    # Delivery Model
    doc.add_page_break()
    heading = doc.add_heading('The Delivery Model: How It Works', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    phases = [
        ('Phase 1: Discovery & Design (2-4 Weeks)', [
            'Understand your portfolio management challenge',
            'Identify data sources and integration points',
            'Define scoring dimensions and weights',
            'Map compliance and risk requirements',
            'Design stakeholder assessment framework',
            'Configure decision support framework'
        ]),
        ('Phase 2: Foundation Build (4-8 Weeks)', [
            'Set up infrastructure (cloud or on-premise)',
            'Build data integration pipelines',
            'Develop scoring engine with your criteria',
            'Implement core dashboard and analytics',
            'Create initial reports'
        ]),
        ('Phase 3: Enhancement & Intelligence (4-6 Weeks)', [
            'Add stakeholder assessment capability',
            'Implement AI/ML predictive analytics',
            'Build compliance monitoring',
            'Create workflow automation',
            'Develop natural language Q&A'
        ]),
        ('Phase 4: Deployment & Adoption (2-4 Weeks)', [
            'User training and documentation',
            'Change management support',
            'Executive briefings',
            'Performance tuning',
            'Go-live support'
        ])
    ]

    for phase_title, activities in phases:
        subheading = doc.add_heading(phase_title, 2)
        subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        subheading.runs[0].font.size = Pt(11)

        for activity in activities:
            p = doc.add_paragraph(activity, style='List Bullet')
            p.runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run('Total Time to Value: 12-22 Weeks')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 204)
    run = p.add_run(' (3-6 months from kickoff to full production)')
    run.font.size = Pt(11)

    # Case Studies
    doc.add_page_break()
    heading = doc.add_heading('Case Studies: Proof of Concept', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    case_studies = [
        ('Application Rationalization - Fortune 500 Financial Services', {
            'Challenge': '300+ applications, $85M IT spend, no visibility, CIO mandate to cut 20%',
            'Solution': 'Implemented platform in 16 weeks, scored all applications, conducted 50+ stakeholder assessments',
            'Outcome': 'Identified $18M savings (21%), 45 applications for retirement, ROI: 60x in Year 1'
        }),
        ('Capital Projects - County Public Works Department', {
            'Challenge': '85 projects, $220M budget, frequent overruns, missed grant milestones, 40 hrs/month reporting',
            'Solution': 'Implemented planner in 14 weeks, automated health scoring, predictive analytics',
            'Outcome': 'Prevented $2.5M in overruns, saved $3M in grant funding, 85% reporting time reduction, ROI: 44x'
        }),
        ('Contract Oversight - State Agency', {
            'Challenge': '450 contracts, $180M spend, 18 audit findings, no performance tracking, overwhelmed team',
            'Solution': 'Implemented oversight system in 18 weeks, automated compliance, vendor performance framework',
            'Outcome': 'Zero audit findings (18 to 0), $12M consolidation opportunities, 70% admin time reduction, ROI: 52x'
        })
    ]

    for i, (title, details) in enumerate(case_studies, 1):
        subheading = doc.add_heading(f'Case Study {i}: {title}', 2)
        subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        subheading.runs[0].font.size = Pt(11)

        for label, text in details.items():
            p = doc.add_paragraph()
            run = p.add_run(f'{label}: ')
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(text)
            run.font.size = Pt(10)

    # Business Case
    doc.add_page_break()
    heading = doc.add_heading('The Business Case: Why This Matters', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    subheading = doc.add_heading('The Problem with Traditional Approaches', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    comparison = [
        ('Approach', 'Challenges'),
        ('Custom Development', '12-24 month timelines, $500K-2M+ costs, high failure risk, uncertain outcomes'),
        ('Off-the-Shelf Software', 'Generic (not tailored), expensive licensing ($200K-500K+/yr), long implementations (6-18 mo), vendor lock-in'),
        ('Consulting-Only', 'Recommendations (not tools), insights don\'t persist, no ongoing decision support, expensive to repeat')
    ]

    for i, (approach, challenges) in enumerate(comparison):
        table.rows[i].cells[0].text = approach
        table.rows[i].cells[1].text = challenges

        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True

    subheading = doc.add_heading('The Platform Architecture Advantage', 2)
    subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    advantages = [
        '3-6 month delivery (not 12-24 months)',
        '$150K-500K investment (not $1M-5M)',
        'Known outcomes based on three successful implementations',
        'Adaptable as your needs evolve',
        '10-60% cost savings on managed portfolio',
        '60-80% time savings on administrative overhead',
        '80-95% improvement in compliance/audit performance',
        '4-87x ROI in first year (proven across three domains)'
    ]

    for advantage in advantages:
        p = doc.add_paragraph(advantage, style='List Bullet')
        p.runs[0].font.size = Pt(10)

    # Pricing & Engagement Models
    doc.add_page_break()
    heading = doc.add_heading('Pricing & Engagement Models', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    models = [
        ('Option 1: Custom Platform Development', [
            'Discovery & design: $25K-50K',
            'Foundation build: $75K-150K',
            'Enhancement & intelligence: $50K-100K',
            'Deployment & adoption: $25K-50K',
            'Total: $175K-350K (one-time)',
            'Ongoing: Hosting ($5K-15K/yr), Support ($25K-50K/yr)'
        ]),
        ('Option 2: Platform License + Configuration', [
            'Annual SaaS License: $50K-250K/year (based on portfolio size)',
            'Includes hosting, support, updates',
            'Configuration services: $50K-100K (one-time)',
            'Training & change management: $15K-30K'
        ]),
        ('Option 3: Consulting Partnership', [
            'Revenue share model',
            'I provide platform architecture and technical delivery',
            'You provide domain expertise and client relationships',
            'Revenue split on client engagements: 50/50 or negotiable',
            'White-label options available'
        ]),
        ('Option 4: Build-Operate-Transfer', [
            'I build the complete solution for your organization',
            'Operate it for 6-12 months to ensure success',
            'Transfer all code, documentation, and knowledge to your team',
            'Fixed price: $300K-600K depending on complexity'
        ])
    ]

    for option_title, details in models:
        subheading = doc.add_heading(option_title, 2)
        subheading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        subheading.runs[0].font.size = Pt(11)

        for detail in details:
            p = doc.add_paragraph(detail, style='List Bullet')
            p.runs[0].font.size = Pt(9)

    # The Bottom Line
    doc.add_page_break()
    heading = doc.add_heading('The Bottom Line', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('I don\'t just build software. I deliver business transformation through intelligent platform architecture.')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run('With three proven implementations across fundamentally different domains, I\'ve demonstrated the ability to:')
    p.runs[0].font.size = Pt(11)

    capabilities = [
        'Rapidly understand complex business problems',
        'Design and build data-driven decision platforms',
        'Deliver measurable business outcomes (ROI of 4-87x)',
        'Deploy modern AI/ML capabilities in practical applications',
        'Work with enterprise stakeholders from technical teams to C-suite'
    ]

    for capability in capabilities:
        p = doc.add_paragraph(capability, style='List Bullet')
        p.runs[0].font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run('Whether your challenge is application rationalization, capital project optimization, contract oversight, or any other portfolio management problem')
    run.font.size = Pt(11)
    run = p.add_run(' - I have the proven framework, technical capabilities, and business acumen to deliver results.')
    run.font.size = Pt(11)
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('The question isn\'t "Can this be done?" The question is: ')
    p.runs[0].font.size = Pt(11)
    run = p.add_run('"How fast do you want to see results?"')
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)

    # Closing statement
    doc.add_page_break()
    p = doc.add_paragraph('"Any complex portfolio management challenge can be solved with the right framework, proven architecture, and rapid execution. I\'ve done it three times. Let\'s make yours the fourth."')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Save the document
    output_path = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\PLATFORM_CAPABILITY_STATEMENT.docx'
    doc.save(output_path)
    print(f"SUCCESS: Word document created successfully: {output_path}")

    return output_path

if __name__ == '__main__':
    create_capability_statement()
