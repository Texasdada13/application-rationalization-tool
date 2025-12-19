"""
Simple, robust markdown to Word converter - processes every line
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re

# Read markdown
with open(r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\PLATFORM_CAPABILITY_STATEMENT.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Create document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Set default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

lines = md_content.split('\n')
i = 0
in_code_block = False
code_lines = []

print(f"Processing {len(lines)} lines...")

while i < len(lines):
    line = lines[i]

    # Code blocks
    if line.strip().startswith('```'):
        if not in_code_block:
            in_code_block = True
            code_lines = []
        else:
            # End code block
            in_code_block = False
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 102, 204)
            code_lines = []
        i += 1
        continue

    if in_code_block:
        code_lines.append(line.strip())
        i += 1
        continue

    # Horizontal rule - skip
    if line.strip() == '---':
        doc.add_paragraph()
        i += 1
        continue

    # Tables - collect all table lines first
    if '|' in line and i + 1 < len(lines) and re.match(r'^[\s\|:\-]+$', lines[i+1]):
        # This is a table
        table_lines = [line]
        i += 1  # Skip separator
        i += 1  # Move to first data row

        # Collect all table rows
        while i < len(lines) and '|' in lines[i] and lines[i].strip():
            table_lines.append(lines[i])
            i += 1

        # Parse table
        headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
        data_rows = [[cell.strip() for cell in row.split('|') if cell.strip()] for row in table_lines[1:]]

        # Create table
        if headers and data_rows:
            table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
            table.style = 'Light Grid Accent 1'

            # Headers
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = re.sub(r'\*\*(.*?)\*\*', r'\1', header)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            # Data
            for row_idx, row_data in enumerate(data_rows):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(headers):
                        cell = table.rows[row_idx + 1].cells[col_idx]
                        cell.text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_data)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)

        doc.add_paragraph()  # Spacing after table
        continue

    # H1
    if line.startswith('# ') and not line.startswith('## '):
        text = line[2:].strip()
        if 'Platform Capability Statement' in text:
            doc.add_paragraph()
            title = doc.add_heading(text.split(':')[0] if ':' in text else text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.size = Pt(32)
                run.font.color.rgb = RGBColor(0, 51, 102)

            if ':' in text:
                subtitle = doc.add_paragraph(text.split(':', 1)[1].strip())
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in subtitle.runs:
                    run.font.size = Pt(20)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 102, 204)

            doc.add_page_break()
        else:
            if i > 10:
                doc.add_page_break()
            h = doc.add_heading(text, 1)
            for run in h.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 51, 102)
        i += 1
        continue

    # H2
    if line.startswith('## ') and not line.startswith('### '):
        h = doc.add_heading(line[3:].strip(), 2)
        for run in h.runs:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 102, 204)
        i += 1
        continue

    # H3
    if line.startswith('### ') and not line.startswith('#### '):
        h = doc.add_heading(line[4:].strip(), 3)
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 102, 204)
        i += 1
        continue

    # H4
    if line.startswith('#### '):
        p = doc.add_paragraph()
        run = p.add_run(line[5:].strip())
        run.bold = True
        run.font.size = Pt(12)
        i += 1
        continue

    # Bullet
    if line.strip().startswith('- '):
        text = line.strip()[2:]
        p = doc.add_paragraph(style='List Bullet')

        if '**' in text:
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                elif part:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
        else:
            run = p.add_run(text)
            run.font.size = Pt(11)
        i += 1
        continue

    # Numbered
    if re.match(r'^\d+\.\s+', line.strip()):
        text = re.sub(r'^\d+\.\s+', '', line.strip())
        p = doc.add_paragraph(style='List Number')

        if '**' in text:
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                elif part:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
        else:
            run = p.add_run(text)
            run.font.size = Pt(11)
        i += 1
        continue

    # Regular paragraph
    if line.strip():
        p = doc.add_paragraph()

        if '**' in line:
            parts = re.split(r'(\*\*.*?\*\*)', line.strip())
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                elif part:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
        else:
            run = p.add_run(line.strip())
            run.font.size = Pt(11)

    i += 1

# Save
output_path = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\PLATFORM_CAPABILITY_STATEMENT.docx'
doc.save(output_path)

print(f"\nSUCCESS: Document saved to {output_path}")
print(f"Final document stats:")
print(f"  - {len(doc.paragraphs)} paragraphs")
print(f"  - {len(doc.tables)} tables")

# Calculate word count
text = ' '.join([p.text for p in doc.paragraphs])
words = len(text.split())
print(f"  - {words:,} words")
print(f"  - Approximately {words/250:.1f} pages")
