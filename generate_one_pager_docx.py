"""
Generate a professionally formatted Word document for the Application Rationalization Tool One-Pager
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)

def create_one_pager():
    """Create the one-pager Word document"""
    doc = Document()

    # Set document margins (narrower for one-page fit)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Title
    title = doc.add_heading('Application Rationalization Tool', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('One-Page Executive Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    # Add spacing
    doc.add_paragraph()

    # What It Is
    heading = doc.add_heading('What It Is', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    p = doc.add_paragraph('A comprehensive platform that transforms application portfolio chaos into actionable insights through automated scoring, stakeholder assessment, and data-driven recommendations.')
    p.runs[0].font.size = Pt(10)

    # The Problem
    heading = doc.add_heading('The Problem', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('Organizations waste 30-40% of IT spend on the wrong applications.')
    run.bold = True
    run.font.size = Pt(10)

    problems = [
        'CIOs manage 200+ applications with no clear view of what to keep, retire, or invest in',
        '$100K-500K+ wasted annually on redundant and underutilized systems',
        'Decisions driven by politics and gut feel instead of objective data',
        'Hidden security vulnerabilities and compliance risks in aging applications',
        'Post-merger duplication and technical debt accumulation'
    ]

    for problem in problems:
        p = doc.add_paragraph(problem, style='List Bullet')
        p.runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run('Bottom line: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Application chaos leads to wasted money, increased risk, and missed strategic opportunities.')
    run.font.size = Pt(9)

    # The Solution
    heading = doc.add_heading('The Solution', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('Data-driven portfolio optimization that combines quantitative metrics with stakeholder insights.')
    run.bold = True
    run.font.size = Pt(10)

    # How It Works
    p = doc.add_paragraph()
    run = p.add_run('How It Works:')
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph('COLLECT → ASSESS → INTERVIEW → ANALYZE → ROADMAP')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    steps = [
        ('Collect', 'application inventory data (costs, owners, usage, tech stack)'),
        ('Score', 'every application on business value and technical health'),
        ('Interview', 'stakeholders using 30+ structured questions across 7 categories'),
        ('Analyze', 'using TIME Framework, ML clustering, compliance checks, and risk assessment'),
        ('Generate', 'prioritized roadmap with specific recommendations for each application')
    ]

    for i, (step_name, step_desc) in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. ', style='List Number')
        run = p.add_run(step_name)
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(f' {step_desc}')
        run.font.size = Pt(9)

    # Key Capabilities Table
    p = doc.add_paragraph()
    run = p.add_run('Key Capabilities')
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'

    capabilities = [
        ('Feature', 'Benefit'),
        ('Automated Scoring', 'Objective, weighted composite scores eliminate bias'),
        ('TIME Framework', 'Categorize every app: Tolerate/Invest/Migrate/Eliminate'),
        ('Stakeholder Assessment', 'Capture tribal knowledge and sentiment systematically'),
        ('Compliance Engine', 'Check against SOX, HIPAA, GDPR, PCI-DSS'),
        ('ML Insights', 'Discover consolidation opportunities through clustering'),
        ('Smart Recommendations', 'AI-generated, prioritized action items'),
        ('Executive Reports', 'PDF, PowerPoint, Excel - ready for board presentation')
    ]

    for i, (feature, benefit) in enumerate(capabilities):
        row = table.rows[i]
        row.cells[0].text = feature
        row.cells[1].text = benefit

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if i == 0:
                        run.bold = True

    # The Outcome
    heading = doc.add_heading('The Outcome', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('Typical Results:')
    run.bold = True
    run.font.size = Pt(9)

    results = [
        '20-30% cost savings through retirement and consolidation',
        'Risk reduction by identifying security and compliance gaps',
        'Strategic clarity on where to invest vs. divest',
        'Stakeholder alignment through shared, objective data',
        'Faster decisions with complete portfolio visibility'
    ]

    for result in results:
        p = doc.add_paragraph(result, style='List Bullet')
        p.runs[0].font.size = Pt(9)

    # Sample ROI
    p = doc.add_paragraph()
    run = p.add_run('Sample ROI: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('150 Applications | $15M Annual Spend')
    run.font.size = Pt(9)

    roi_items = [
        'Retire 20 applications → $1.2M saved',
        'Consolidate 15 applications → $600K saved',
        'Optimize licensing → $300K saved',
        'Total: $2.5M savings (17% reduction)',
        'Investment: $300K-600K',
        'ROI: 4-8x in Year 1'
    ]

    for item in roi_items:
        p = doc.add_paragraph(f'• {item}')
        p.runs[0].font.size = Pt(9)
        if 'Total:' in item or 'ROI:' in item:
            p.runs[0].font.bold = True

    # Two column section: Who It's For and What Makes It Different
    table = doc.add_table(rows=1, cols=2)

    # Who It's For (left column)
    left_cell = table.rows[0].cells[0]
    p = left_cell.paragraphs[0]
    run = p.add_run('Who It\'s For')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 102)

    audiences = [
        'CIOs/CTOs optimizing IT spend',
        'IT Portfolio Managers',
        'M&A Teams',
        'Digital Transformation Leaders',
        'Consultants'
    ]

    for audience in audiences:
        p = left_cell.add_paragraph(audience, style='List Bullet')
        p.runs[0].font.size = Pt(8)

    # What Makes It Different (right column)
    right_cell = table.rows[0].cells[1]
    p = right_cell.paragraphs[0]
    run = p.add_run('What Makes It Different')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. Spreadsheets: ')
    run.font.size = Pt(8)
    run.bold = True
    run = p.add_run('Automated, historical tracking, ML analysis')
    run.font.size = Pt(8)

    p = right_cell.add_paragraph()
    run = p.add_run('vs. Enterprise Tools: ')
    run.font.size = Pt(8)
    run.bold = True
    run = p.add_run('Days to deploy, affordable, focused')
    run.font.size = Pt(8)

    p = right_cell.add_paragraph()
    run = p.add_run('Unique Strengths:')
    run.font.size = Pt(8)
    run.bold = True

    strengths = [
        'Stakeholder assessment built-in',
        'Interview-to-score pipeline',
        'Action-oriented design',
        'Consultant-ready'
    ]

    for strength in strengths:
        p = right_cell.add_paragraph(strength, style='List Bullet')
        p.runs[0].font.size = Pt(8)

    # Use Cases Table
    heading = doc.add_heading('Proven Use Cases', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light List Accent 1'

    use_cases = [
        ('Scenario', 'Challenge', 'Solution', 'Result'),
        ('Post-Merger', 'Duplicate systems', 'Assess & identify overlaps', '40% reduction, $3M saved'),
        ('Cloud Migration', 'Which apps first?', 'Score & prioritize', 'Phased 50-app plan'),
        ('Budget Defense', 'CFO demands 15% cut', 'Show impact data', '12% savings, protected strategic apps'),
        ('Technical Debt', 'Team overwhelmed', 'Prioritize retirement', 'Retired 15 apps, freed 3 FTEs'),
        ('Compliance Audit', 'SOX audit prep', 'Run compliance check', 'Zero findings, fixed 23 gaps')
    ]

    for i, (scenario, challenge, solution, result) in enumerate(use_cases):
        row = table.rows[i]
        row.cells[0].text = scenario
        row.cells[1].text = challenge
        row.cells[2].text = solution
        row.cells[3].text = result

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
                    if i == 0:
                        run.bold = True

    # The Bottom Line
    heading = doc.add_heading('The Bottom Line', 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run('Turn application chaos into confident, data-driven decisions.')
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph('Most organizations know they\'re wasting money on applications - they just don\'t know where to start. This platform provides the objective data, stakeholder insights, and strategic framework to make portfolio decisions with confidence.')
    p.runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run('Stop investing in the wrong applications. Start optimizing your portfolio.')
    run.font.size = Pt(9)
    run.italic = True

    # Get Started
    p = doc.add_paragraph()
    run = p.add_run('8-12 Week Typical Engagement: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run('Setup → Interviews → Analysis → Recommendations → Roadmap')
    run.font.size = Pt(9)

    # Footer tagline
    doc.add_paragraph()
    p = doc.add_paragraph('"We help you see your application portfolio clearly, so you can invest wisely and cut confidently."')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Save the document
    output_path = r'C:\Users\dada_\OneDrive\Documents\application-rationalization-tool\docs\ONE_PAGER.docx'
    doc.save(output_path)
    print(f"SUCCESS: Word document created successfully: {output_path}")

    return output_path

if __name__ == '__main__':
    create_one_pager()
